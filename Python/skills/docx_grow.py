"""
docx "成长"编辑 skill
-----------------------

让 AI 可以在已有 docx 上精准叠加内容，不重写整个文档：

- insert_section_after_heading: 在指定标题之后插入一个新章节（含标题 + 段落）
- append_paragraphs_after_heading: 在指定标题下追加几段正文
- append_table_row: 向第 N 个表格追加一行
- insert_bullets_after_heading: 追加无序列表

所有操作都在目标文件副本上做，返回新文件路径，避免污染原稿。
用户可以选择"确认采纳"覆盖原文件（UI 层做）。

关键实现细节：
python-docx 没有 "在某段之后插入" 的一等 API，但 Paragraph._p 可以用
`_p.addprevious()` / `_p.addnext()` 挂 lxml element，我们用这条路。
"""

from __future__ import annotations

import copy
import shutil
from dataclasses import dataclass
from pathlib import Path
from typing import Any


class DocxGrowError(RuntimeError):
    pass


@dataclass
class GrowResult:
    output_path: str
    inserted_elements: int
    notes: list[str]


def _open_copy(source_path: str, output_path: str | None) -> tuple["object", Path]:
    """python-docx 加载目标副本。src 和 out 是同一文件时原地打开。"""
    try:
        from docx import Document
    except ImportError as exc:
        raise DocxGrowError("需要 python-docx；venv 里没装上") from exc

    src = Path(source_path).expanduser().resolve()
    if not src.exists():
        raise DocxGrowError(f"文件不存在：{src}")

    if output_path:
        out = Path(output_path).expanduser().resolve()
    else:
        out = src.with_name(src.stem + "-edited.docx")

    out.parent.mkdir(parents=True, exist_ok=True)
    if out != src:
        shutil.copyfile(src, out)
    return Document(str(out)), out


def _heading_level_of(paragraph: "object") -> int | None:
    """0 = 非标题；1/2/3... = Heading N；中文本地化"标题 N"也认。"""
    style_name = getattr(paragraph.style, "name", "") or ""
    for prefix in ("Heading ", "标题 "):
        if style_name.startswith(prefix):
            try:
                return int(style_name[len(prefix):].strip())
            except ValueError:
                return None
    return None


def _find_section_end_anchor(doc: "object", heading_para: "object") -> "object":
    """从 heading_para 开始往后扫，找到同级或更高级的下一个标题。
    该节内容（正文 + 表格）的末尾 anchor 就是这个"下一标题的上一个 element"。
    返回应该 addnext 的 anchor（插入点会出现在它之后）。
    """
    heading_level = _heading_level_of(heading_para) or 99

    # 在 body 的 element 流里定位 heading_para._p
    body = doc.element.body
    children = list(body.iterchildren())
    try:
        start_idx = children.index(heading_para._p)
    except ValueError:
        return heading_para  # 找不到就退回

    # 建立 paragraph._p → paragraph 的映射，便于判断 heading
    para_by_p = {p._p: p for p in doc.paragraphs}

    W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    for i in range(start_idx + 1, len(children)):
        child = children[i]
        if child.tag == f"{{{W_NS}}}p":
            p = para_by_p.get(child)
            if p is None:
                continue
            lvl = _heading_level_of(p)
            if lvl is not None and lvl <= heading_level:
                # 插在上一个 element 的 next（= 紧挨当前 heading 的前面）
                # 这里我们返回的 anchor 就是 children[i-1]，调用方对它 addnext
                # 包装成一个类，.addnext() 挂 element 用
                return _PElementWrapper(children[i - 1])

    # 没找到下一个同级 heading → 插在文档末尾，anchor = 最后一个 element
    if children:
        return _PElementWrapper(children[-1])
    return heading_para


class _PElementWrapper:
    """给 lxml element 包一层，让 `_make_paragraph_after` 能一视同仁地用 `._p.addnext`。"""
    def __init__(self, element: object) -> None:
        self._p = element


def _find_heading_by_text(doc: "object", heading_text: str, level: int | None = None) -> Any:
    """在 doc.paragraphs 里按文本精确匹配找段落；level 限制到某一级别。
    返回段落对象，找不到返回 None。
    """
    for p in doc.paragraphs:
        if (p.text or "").strip() == heading_text.strip():
            if level is not None:
                style_name = getattr(p.style, "name", "") or ""
                if not style_name.startswith(f"Heading {level}"):
                    continue
            return p
    return None


def _make_paragraph_after(
    anchor: "object",
    text: str,
    style: str | None = None,
    doc: "object | None" = None,
) -> None:
    """在 anchor 段落之后插入一个新段落。

    如果传了 style 且 doc 可用：借用 python-docx 的 add_paragraph 先造再 lxml 搬家，
    这样 style 的 pStyle/@w:val 会用对 style_id（"Heading 2" → "Heading2"）。
    """
    if style and doc is not None:
        try:
            tmp_p = doc.add_paragraph(text or "", style=style)
            p_element = tmp_p._p
            anchor._p.addnext(p_element)
            return
        except KeyError:
            # style 名字在文档里不存在，回退到手写 XML
            pass

    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    new_p = OxmlElement("w:p")
    if style:
        pPr = OxmlElement("w:pPr")
        pStyle = OxmlElement("w:pStyle")
        # 尝试把"Heading 2"→"Heading2"
        pStyle.set(qn("w:val"), style.replace(" ", ""))
        pPr.append(pStyle)
        new_p.append(pPr)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text or ""
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    new_p.append(r)
    anchor._p.addnext(new_p)


# ---- 公开 API ----


def insert_section_after_heading(
    source_path: str,
    *,
    after_heading: str,
    new_heading: str,
    new_heading_level: int = 2,
    paragraphs: list[str] | None = None,
    anchor_level: int | None = None,
    output_path: str | None = None,
) -> GrowResult:
    """在 after_heading 之后插入一个新章节：新标题 + 若干段落。

    插入位置是在 after_heading 所在段落之后、同级下一个同级别标题之前。
    （新标题和段落物理上紧跟 anchor 段落）
    """
    doc, out = _open_copy(source_path, output_path)
    anchor_heading = _find_heading_by_text(doc, after_heading, level=anchor_level)
    if anchor_heading is None:
        raise DocxGrowError(f"找不到标题：{after_heading!r}")

    # 真实插入锚点：anchor_heading 所属章节的末尾
    section_end = _find_section_end_anchor(doc, anchor_heading)

    notes: list[str] = []
    count = 0
    style = f"Heading {new_heading_level}"

    # 倒序 addnext：最先 addnext 的会被后面的 element 推到后面
    # 想要最终顺序 [new_heading, p1, p2, p3]，倒着 addnext 依次 p3, p2, p1, new_heading
    if paragraphs:
        for text in reversed(paragraphs):
            _make_paragraph_after(section_end, text, doc=doc)
            count += 1
    _make_paragraph_after(section_end, new_heading, style=style, doc=doc)
    count += 1
    notes.append(f"在 {after_heading!r} 后的章节末尾插入 {new_heading!r}（{count} 个 element）")

    doc.save(str(out))
    return GrowResult(output_path=str(out), inserted_elements=count, notes=notes)


def append_paragraphs_after_heading(
    source_path: str,
    *,
    after_heading: str,
    paragraphs: list[str],
    anchor_level: int | None = None,
    output_path: str | None = None,
) -> GrowResult:
    """在 after_heading 所在章节的末尾追加若干段落（不新建子标题）。"""
    doc, out = _open_copy(source_path, output_path)
    anchor = _find_heading_by_text(doc, after_heading, level=anchor_level)
    if anchor is None:
        raise DocxGrowError(f"找不到标题：{after_heading!r}")
    section_end = _find_section_end_anchor(doc, anchor)

    count = 0
    for text in reversed(paragraphs):
        _make_paragraph_after(section_end, text)
        count += 1

    doc.save(str(out))
    return GrowResult(
        output_path=str(out),
        inserted_elements=count,
        notes=[f"在 {after_heading!r} 章节末尾追加 {count} 段"],
    )


def insert_bullets_after_heading(
    source_path: str,
    *,
    after_heading: str,
    bullets: list[str],
    anchor_level: int | None = None,
    output_path: str | None = None,
    bullet_style: str = "List Bullet",
) -> GrowResult:
    """在 after_heading 章节末尾追加一组无序列表项。style 需模板里已定义。"""
    doc, out = _open_copy(source_path, output_path)
    anchor = _find_heading_by_text(doc, after_heading, level=anchor_level)
    if anchor is None:
        raise DocxGrowError(f"找不到标题：{after_heading!r}")
    section_end = _find_section_end_anchor(doc, anchor)

    count = 0
    for text in reversed(bullets):
        _make_paragraph_after(section_end, text, style=bullet_style, doc=doc)
        count += 1

    doc.save(str(out))
    return GrowResult(
        output_path=str(out),
        inserted_elements=count,
        notes=[f"追加 {count} 个无序列表项"],
    )


def append_table_row(
    source_path: str,
    *,
    table_index: int,
    cells: list[str],
    output_path: str | None = None,
) -> GrowResult:
    """向 doc.tables[table_index] 追加一行。单元格数要与表格列数一致，
    少了用 "" 补，多了截断。
    """
    doc, out = _open_copy(source_path, output_path)
    if table_index < 0 or table_index >= len(doc.tables):
        raise DocxGrowError(
            f"table_index={table_index} 超出范围（共 {len(doc.tables)} 个表）"
        )
    tbl = doc.tables[table_index]
    n_cols = len(tbl.columns)
    padded = (cells + [""] * n_cols)[:n_cols]
    new_row = tbl.add_row()
    for idx, text in enumerate(padded):
        new_row.cells[idx].text = text or ""

    doc.save(str(out))
    return GrowResult(
        output_path=str(out),
        inserted_elements=1,
        notes=[f"表格 #{table_index} 追加一行：{padded}"],
    )


def list_headings(source_path: str) -> list[dict[str, Any]]:
    """列出文档里所有标题（Heading 1/2/3...），用于 AI 定位锚点。

    返回 [{"level": int, "text": str, "index": int}, ...]
    """
    try:
        from docx import Document
    except ImportError as exc:
        raise DocxGrowError("需要 python-docx") from exc

    doc = Document(str(Path(source_path).expanduser().resolve()))
    out: list[dict[str, Any]] = []
    for i, p in enumerate(doc.paragraphs):
        style_name = getattr(p.style, "name", "") or ""
        if not style_name.startswith("Heading "):
            continue
        try:
            level = int(style_name.split(" ")[1])
        except (ValueError, IndexError):
            level = 0
        text = (p.text or "").strip()
        if text:
            out.append({"level": level, "text": text, "index": i})
    return out
